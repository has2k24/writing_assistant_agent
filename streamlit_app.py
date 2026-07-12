"""
Inkwell — AI Writing Assistant (Streamlit)
IBM Watsonx.ai · Llama 3.3 70B Instruct
"""

import os, re, io, datetime
import streamlit as st
from dotenv import load_dotenv
from ibm_watsonx_ai import APIClient, Credentials
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
#  AGENT INSTRUCTIONS — Customise freely below
# ─────────────────────────────────────────────
AGENT_INSTRUCTIONS = {
    "persona": (
        "You are Inkwell, a warm, knowledgeable AI writing assistant. "
        "You ONLY answer questions and assist with what the user explicitly asks. "
        "You never volunteer to write content on behalf of the user unless directly asked. "
        "You guide, advise, define, explain, and provide feedback — you do not write for the user."
    ),
    "tone": (
        "Be calm, clear, and encouraging. Never be preachy or over-enthusiastic. "
        "Give direct, actionable answers. Keep responses focused and concise."
    ),
    "specializations": [
        "Grammar checking and correction",
        "Vocabulary — definitions, synonyms, pronunciations",
        "Writing advice and technique guidance",
        "Structural feedback on drafts",
        "Indian regional literature styles (Hindi, Tamil, Telugu, Kannada, Bengali aesthetics)",
    ],
    "indian_culture_palette": (
        "You may use Indian cultural references (chai, biryani, monsoon, Diwali, Ramayana, "
        "Carnatic music, Bharatanatyam) as illustrative examples when relevant and natural."
    ),
    "grammar_rules": (
        "When checking grammar: list each error clearly as "
        "❌ [wrong phrase] → ✅ [corrected phrase] — Reason: [plain English explanation]. "
        "If no errors exist, say so explicitly."
    ),
    "word_features": (
        "For definitions: give part of speech, clear definition, one example sentence, "
        "phonetic pronunciation like 'Pronounced as: ser-EN-i-tee', and brief etymology. "
        "For synonyms: list 6-8 with a one-line meaning each. "
        "For sentence descriptors: give mood, style, imagery, and 5 descriptor words."
    ),
    "safety_rules": (
        "Never produce hateful, explicit, or harmful content. "
        "Decline politely if asked to write something inappropriate."
    ),
    "response_format": (
        "Use markdown for structure where helpful. "
        "Grammar corrections use: ❌ wrong → ✅ correct — Reason: explanation. "
        "Keep answers focused. Do not pad with unnecessary preamble."
    ),
}

# ─────────────────────────────────────────────
#  Load credentials
# ─────────────────────────────────────────────
_base = os.path.dirname(os.path.abspath(__file__))
for _f in [os.path.join(_base, ".env"), os.path.join(_base, ".env.example")]:
    if os.path.exists(_f):
        load_dotenv(_f)
        break

IBM_API_KEY    = os.getenv("IBM_API_KEY", "")
IBM_PROJECT_ID = os.getenv("IBM_PROJECT_ID", "")
IBM_REGION     = os.getenv("IBM_REGION", "us-south")
WATSONX_URL    = f"https://{IBM_REGION}.ml.cloud.ibm.com"
MODEL_ID       = "meta-llama/llama-3-3-70b-instruct"

# ─────────────────────────────────────────────
#  Watsonx model — cached singleton
# ─────────────────────────────────────────────
@st.cache_resource(show_spinner="Connecting to IBM Watsonx.ai…")
def get_model():
    creds  = Credentials(url=WATSONX_URL, api_key=IBM_API_KEY)
    client = APIClient(creds)
    return ModelInference(
        model_id=MODEL_ID,
        api_client=client,
        project_id=IBM_PROJECT_ID,
        params={
            GenParams.MAX_NEW_TOKENS:      1024,
            GenParams.TEMPERATURE:         0.7,
            GenParams.TOP_P:               0.95,
            GenParams.REPETITION_PENALTY:  1.1,
        },
    )

def build_system_prompt() -> str:
    ai = AGENT_INSTRUCTIONS
    return "\n".join([
        ai["persona"],
        "Tone: "              + ai["tone"],
        "Specializations: "   + ", ".join(ai["specializations"]),
        "Cultural palette: "  + ai["indian_culture_palette"],
        "Grammar guidance: "  + ai["grammar_rules"],
        "Word features: "     + ai["word_features"],
        "Safety: "            + ai["safety_rules"],
        "Response format: "   + ai["response_format"],
    ])

def call_model(prompt: str, max_tokens: int = 900, temperature: float = 0.7) -> str:
    model  = get_model()
    system = build_system_prompt()
    full   = f"<|system|>\n{system}\n<|user|>\n{prompt}\n<|assistant|>\n"
    model.params[GenParams.MAX_NEW_TOKENS] = max_tokens
    model.params[GenParams.TEMPERATURE]    = temperature
    return model.generate_text(prompt=full).strip()

# ─────────────────────────────────────────────
#  DOCX export helper
# ─────────────────────────────────────────────
def build_docx(title: str, content: str, author: str) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Garamond"
    doc.styles["Normal"].font.size = Pt(12)
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"{author}  |  {datetime.date.today().strftime('%B %d, %Y')}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    mr.italic = True
    doc.add_paragraph()
    for line in content.split("\n"):
        p = doc.add_paragraph(line.strip()) if line.strip() else doc.add_paragraph()
        if line.strip():
            p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— Generated with Inkwell AI Writing Assistant —")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
    fr.italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────
#  Page config + custom CSS
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Inkwell — Writing Assistant",
    page_icon="🖋",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
/* ── Font import ── */
@import url('https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,600;1,400&family=Inter:wght@300;400;500;600&display=swap');

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'Inter', system-ui, sans-serif;
}

/* ── Hide default Streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding-top: 1.2rem !important; padding-bottom: 1rem !important; }

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: #13161e !important;
    border-right: 1px solid #1f2433;
}
section[data-testid="stSidebar"] * { color: #c8cdd8 !important; }
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stTextInput label,
section[data-testid="stSidebar"] .stTextArea label { color: #8a93a8 !important; font-size: 0.75rem !important; }

/* ── Chat messages ── */
.msg-user {
    background: #2d2850;
    border: 1px solid #7c6af0;
    border-radius: 14px 14px 4px 14px;
    padding: 0.7rem 1rem;
    margin: 0.4rem 0 0.4rem 15%;
    font-size: 0.88rem;
    line-height: 1.65;
    color: #e8ecf4;
}
.msg-agent {
    background: #1a1e2a;
    border: 1px solid #2a2f3d;
    border-radius: 14px 14px 14px 4px;
    padding: 0.7rem 1rem;
    margin: 0.4rem 15% 0.4rem 0;
    font-size: 0.88rem;
    line-height: 1.65;
    color: #e8ecf4;
}
.msg-meta {
    font-size: 0.65rem;
    color: #5a6278;
    margin-top: 3px;
    padding: 0 4px;
}

/* ── Draft textarea ── */
.draft-area textarea {
    font-family: 'Lora', Georgia, serif !important;
    font-size: 1rem !important;
    line-height: 1.9 !important;
    background: #0d0f14 !important;
    color: #e8ecf4 !important;
    border: 1px solid #2a2f3d !important;
    border-radius: 10px !important;
    padding: 1.2rem 1.5rem !important;
}

/* ── Result cards ── */
.result-card {
    background: #1a1e2a;
    border: 1px solid #2a2f3d;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.6rem;
    font-size: 0.85rem;
    line-height: 1.75;
    color: #c8cdd8;
}
.result-card strong { color: #a78bfa; }
.result-card em { color: #d4a853; font-style: italic; }
.result-card .label {
    font-size: 0.65rem; font-weight: 600; letter-spacing: 0.1em;
    text-transform: uppercase; color: #7c6af0; margin-bottom: 0.4rem;
}
.word-title {
    font-family: 'Lora', Georgia, serif;
    font-size: 1.3rem; font-weight: 600;
    color: #e8ecf4; margin-bottom: 0.2rem;
}
.pos { display:inline-block; font-size:0.65rem; padding:1px 8px;
       border-radius:20px; background:#2e2414; color:#d4a853;
       border:1px solid rgba(212,168,83,.35); margin-bottom:0.45rem; }
.pron {
    border-left: 3px solid #7c6af0; padding-left: 8px;
    font-style: italic; color: #8a93a8; font-size: 0.8rem; margin: 0.4rem 0;
}
.grammar-err  { color:#f87171; background:#2a0f0f; padding:1px 6px; border-radius:4px; font-family:monospace; font-size:0.82rem; }
.grammar-ok   { color:#4ade80; background:#0f2a1a; padding:1px 6px; border-radius:4px; font-family:monospace; font-size:0.82rem; }
.grammar-reason { color:#8a93a8; font-size:0.78rem; }

/* ── Brand header ── */
.brand-header {
    display: flex; align-items: center; gap: 10px;
    padding: 0.5rem 0 1rem 0;
}
.brand-icon {
    width: 34px; height: 34px;
    background: linear-gradient(135deg, #7c6af0, #a78bfa);
    border-radius: 9px; display: flex; align-items: center;
    justify-content: center; font-size: 1rem;
    box-shadow: 0 0 16px rgba(124,106,240,0.3);
}
.brand-name {
    font-family: 'Lora', Georgia, serif;
    font-size: 1.25rem; font-weight: 600; color: #e8ecf4;
    letter-spacing: -0.01em;
}
.brand-sub { font-size: 0.72rem; color: #5a6278; }

/* ── Tab styling ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px;
    background: #13161e;
    border-radius: 10px;
    padding: 4px;
    border: 1px solid #2a2f3d;
}
.stTabs [data-baseweb="tab"] {
    background: transparent;
    border-radius: 7px;
    color: #8a93a8;
    font-size: 0.8rem;
    font-weight: 500;
    padding: 6px 16px;
    border: none;
}
.stTabs [aria-selected="true"] {
    background: #1e2330 !important;
    color: #e8ecf4 !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.3);
}
.stTabs [data-baseweb="tab-panel"] { padding-top: 1rem; }

/* ── Buttons ── */
.stButton > button {
    background: #1e2330;
    border: 1px solid #2a2f3d;
    color: #c8cdd8;
    border-radius: 8px;
    font-size: 0.8rem;
    font-weight: 500;
    padding: 0.4rem 1rem;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: #7c6af0;
    color: #a78bfa;
    background: #2d2850;
}

/* ── Inputs ── */
.stTextInput > div > div > input,
.stSelectbox > div > div,
.stTextArea > div > div > textarea {
    background: #1a1e2a !important;
    border: 1px solid #2a2f3d !important;
    color: #e8ecf4 !important;
    border-radius: 8px !important;
    font-size: 0.85rem !important;
}

/* ── Section divider ── */
hr { border-color: #1f2433 !important; margin: 0.75rem 0 !important; }

/* ── Word count badge ── */
.wc-badge {
    display: inline-block;
    font-size: 0.68rem; color: #5a6278;
    background: #1a1e2a; border: 1px solid #2a2f3d;
    border-radius: 20px; padding: 2px 10px;
    font-family: monospace;
}

/* ── Spotify widget ── */
.spotify-box {
    background: #111;
    border: 1px solid #1db95433;
    border-radius: 10px;
    padding: 0.75rem;
    margin-top: 0.5rem;
}
.spotify-box a {
    display: block; width: 100%;
    background: #1db954; color: #fff !important;
    text-align: center; border-radius: 7px;
    padding: 0.45rem; font-size: 0.8rem;
    font-weight: 600; text-decoration: none;
}
.spotify-box a:hover { background: #1ed760; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  Session state initialisation
# ─────────────────────────────────────────────
def _init(key, val):
    if key not in st.session_state:
        st.session_state[key] = val

_init("chat_history", [])   # list of {"role": "user"|"assistant", "content": str}
_init("draft_text", "")
_init("draft_title", "")
_init("word_result", "")
_init("grammar_result", "")
_init("active_tool_tab", 0)

# ─────────────────────────────────────────────
#  SIDEBAR — Brand + Tools nav + Spotify
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="brand-header">
      <div class="brand-icon">🖋</div>
      <div>
        <div class="brand-name">Inkwell</div>
        <div class="brand-sub">AI Writing Assistant</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("##### 🧭 Navigate")
    nav = st.radio(
        "", ["💬 Chat", "✍ Draft & Grammar", "📖 Word Tools"],
        label_visibility="collapsed",
        key="nav_radio",
    )

    st.markdown("---")
    st.markdown("##### ⚙️ Model")
    st.caption(f"**Model:** `llama-3-3-70b-instruct`")
    st.caption(f"**Region:** `{IBM_REGION}`")
    st.caption(f"**Project:** `{IBM_PROJECT_ID[:8]}…`" if IBM_PROJECT_ID else "⚠️ Project ID missing")

    st.markdown("---")
    st.markdown("##### 🎵 Spotify")
    st.markdown("""
    <div class="spotify-box">
      <a href="https://open.spotify.com" target="_blank">▶ Open Spotify Web Player</a>
      <div style="font-size:0.68rem;color:#5a6278;margin-top:6px;text-align:center;">
        Opens in a new tab — play music while you write
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    with st.expander("🎨 Quick Prompts", expanded=False):
        prompts = [
            "What does 'in medias res' mean?",
            "What's the difference between show and tell?",
            "How do I write a strong opening sentence?",
            "Explain active vs passive voice with examples.",
            "What makes a metaphor effective?",
            "How do I fix a run-on sentence?",
            "What is the rule for comma splices?",
            "Explain the three-act structure briefly.",
        ]
        for p in prompts:
            if st.button(p, key=f"qp_{p[:20]}", use_container_width=True):
                st.session_state.chat_history.append({"role": "user", "content": p})
                with st.spinner("Thinking…"):
                    try:
                        reply = call_model(p, max_tokens=600)
                        st.session_state.chat_history.append({"role": "assistant", "content": reply})
                    except Exception as e:
                        st.session_state.chat_history.append({"role": "assistant", "content": f"❌ {e}"})
                st.session_state.nav_radio = "💬 Chat"
                st.rerun()

# ─────────────────────────────────────────────
#  MAIN AREA
# ─────────────────────────────────────────────

# ══════════════════════════════════════════════
#  TAB 1 — CHAT
# ══════════════════════════════════════════════
if nav == "💬 Chat":
    st.markdown("### 💬 Ask Inkwell")
    st.caption("Ask about grammar, vocabulary, writing technique, definitions — anything writing-related.")

    # Render conversation
    chat_container = st.container(height=480)
    with chat_container:
        if not st.session_state.chat_history:
            st.markdown("""
            <div class="msg-agent">
                Namaste 🙏 I'm <strong>Inkwell</strong>, your writing assistant.<br><br>
                Ask me about <strong>grammar rules</strong>, <strong>word meanings</strong>,
                <strong>writing techniques</strong>, <strong>synonyms</strong>,
                <strong>sentence structure</strong> — anything to do with writing.<br><br>
                I'm here to assist, not to write for you.
            </div>
            <div class="msg-meta">Inkwell · now</div>
            """, unsafe_allow_html=True)
        else:
            for msg in st.session_state.chat_history:
                t = datetime.datetime.now().strftime("%H:%M")
                if msg["role"] == "user":
                    st.markdown(f'<div class="msg-user">{msg["content"]}</div><div class="msg-meta" style="text-align:right">You · {t}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="msg-agent">{msg["content"]}</div><div class="msg-meta">Inkwell · {t}</div>', unsafe_allow_html=True)

    # Input row
    with st.form("chat_form", clear_on_submit=True):
        col_in, col_btn = st.columns([5, 1])
        with col_in:
            user_msg = st.text_input(
                "Your question",
                placeholder="e.g. What is the Oxford comma? / Define 'ephemeral' / Is my sentence grammatically correct?",
                label_visibility="collapsed",
            )
        with col_btn:
            submitted = st.form_submit_button("Send ➤", use_container_width=True)

    if submitted and user_msg.strip():
        st.session_state.chat_history.append({"role": "user", "content": user_msg.strip()})
        # Build context from last 6 turns
        ctx = ""
        for turn in st.session_state.chat_history[-6:]:
            role = "User" if turn["role"] == "user" else "Inkwell"
            ctx += f"{role}: {turn['content']}\n"
        prompt = (
            "The following is a conversation. Answer ONLY the user's question. "
            "Do not volunteer to write content for them.\n\n"
            f"{ctx}\nInkwell:"
        )
        with st.spinner("Thinking…"):
            try:
                reply = call_model(prompt, max_tokens=700, temperature=0.65)
                reply = re.sub(r"^Inkwell:\s*", "", reply, flags=re.IGNORECASE).strip()
            except Exception as e:
                reply = f"❌ Error: {e}"
        st.session_state.chat_history.append({"role": "assistant", "content": reply})
        st.rerun()

    if st.session_state.chat_history:
        if st.button("🗑 Clear conversation", key="clear_chat"):
            st.session_state.chat_history = []
            st.rerun()

# ══════════════════════════════════════════════
#  TAB 2 — DRAFT & GRAMMAR
# ══════════════════════════════════════════════
elif nav == "✍ Draft & Grammar":
    st.markdown("### ✍ Your Draft")
    st.caption("Write your own content here. Use the tools below to check and fix grammar.")

    # Title
    st.session_state.draft_title = st.text_input(
        "Draft title", value=st.session_state.draft_title,
        placeholder="Untitled draft…",
    )

    # Word count
    wc = len(st.session_state.draft_text.split()) if st.session_state.draft_text.strip() else 0
    cc = len(st.session_state.draft_text)
    st.markdown(f'<span class="wc-badge">📝 {wc} words · {cc} characters</span>', unsafe_allow_html=True)

    # Draft textarea
    st.markdown('<div class="draft-area">', unsafe_allow_html=True)
    st.session_state.draft_text = st.text_area(
        "Draft content",
        value=st.session_state.draft_text,
        height=340,
        placeholder="Write your draft here…",
        label_visibility="collapsed",
    )
    st.markdown("</div>", unsafe_allow_html=True)

    # Grammar action buttons
    col1, col2, col3, col4 = st.columns([2, 2, 2, 2])

    # ── Check Grammar ──
    with col1:
        if st.button("🔍 Check Grammar", use_container_width=True):
            if not st.session_state.draft_text.strip():
                st.warning("Write something first.")
            else:
                with st.spinner("Checking grammar…"):
                    try:
                        prompt = (
                            "You are a professional grammar checker.\n\n"
                            "Check the following text for grammar, spelling, and punctuation errors.\n"
                            "For each error, respond on a new line exactly as:\n"
                            "❌ [wrong phrase] → ✅ [corrected phrase] — Reason: [plain English explanation]\n\n"
                            "If there are NO errors, write exactly: NO_ERRORS\n\n"
                            "After all errors, on a new line write:\n"
                            "CORRECTED_TEXT: [the fully corrected version]\n\n"
                            f"Text:\n{st.session_state.draft_text}"
                        )
                        st.session_state.grammar_result = call_model(prompt, max_tokens=1400, temperature=0.2)
                    except Exception as e:
                        st.session_state.grammar_result = f"ERROR: {e}"

    # ── Fix Grammar ──
    with col2:
        if st.button("✨ Fix Grammar", use_container_width=True):
            if not st.session_state.draft_text.strip():
                st.warning("Write something first.")
            else:
                with st.spinner("Fixing grammar…"):
                    try:
                        prompt = (
                            "Fix ALL grammar, spelling, and punctuation errors in this text. "
                            "Return ONLY the corrected text — no labels, no explanation.\n\n"
                            f"Text:\n{st.session_state.draft_text}"
                        )
                        fixed = call_model(prompt, max_tokens=2000, temperature=0.15)
                        st.session_state.draft_text = fixed
                        st.session_state.grammar_result = ""
                        st.success("Grammar fixed and applied to draft.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Fix failed: {e}")

    # ── Copy button ──
    with col3:
        st.download_button(
            "📋 Copy text",
            data=st.session_state.draft_text,
            file_name="draft.txt",
            mime="text/plain",
            use_container_width=True,
        )

    # ── Clear draft ──
    with col4:
        if st.button("🗑 Clear draft", use_container_width=True):
            st.session_state.draft_text = ""
            st.session_state.draft_title = ""
            st.session_state.grammar_result = ""
            st.rerun()

    # Grammar results
    if st.session_state.grammar_result:
        st.markdown("---")
        st.markdown("##### Grammar Results")
        raw = st.session_state.grammar_result

        if "NO_ERRORS" in raw:
            st.success("✅ No errors found — your text looks great!")
        elif raw.startswith("ERROR:"):
            st.error(raw)
        else:
            lines = raw.split("\n")
            corrected = ""
            error_lines = []
            for line in lines:
                if line.startswith("CORRECTED_TEXT:"):
                    corrected = line.replace("CORRECTED_TEXT:", "").strip()
                elif line.strip():
                    error_lines.append(line)

            if error_lines:
                for line in error_lines:
                    # Try to parse structured error lines
                    m = re.match(r"❌\s*(.+?)\s*→\s*✅\s*(.+?)(?:\s*—\s*Reason:\s*(.*))?$", line)
                    if m:
                        wrong, right, reason = m.group(1).strip(), m.group(2).strip(), (m.group(3) or "").strip()
                        st.markdown(
                            f'<div class="result-card">'
                            f'<span class="grammar-err">✗ {wrong}</span> → <span class="grammar-ok">✓ {right}</span>'
                            + (f'<br><span class="grammar-reason">Reason: {reason}</span>' if reason else "")
                            + '</div>',
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(f'<div class="result-card">{line}</div>', unsafe_allow_html=True)

            if corrected:
                with st.expander("📄 View corrected version"):
                    st.text_area("Corrected text", corrected, height=200, key="corrected_view")
                    if st.button("Apply corrected version to draft"):
                        st.session_state.draft_text = corrected
                        st.session_state.grammar_result = ""
                        st.rerun()

        if st.button("✕ Dismiss results", key="dismiss_grammar"):
            st.session_state.grammar_result = ""
            st.rerun()

    # Download as DOCX
    st.markdown("---")
    with st.expander("📄 Download draft as Word document"):
        dl_author = st.text_input("Author name", placeholder="Your name", key="dl_author")
        if st.session_state.draft_text.strip():
            docx_bytes = build_docx(
                title   = st.session_state.draft_title or "My Draft",
                content = st.session_state.draft_text,
                author  = dl_author or "Written with Inkwell",
            )
            safe_name = re.sub(r"[^\w\s-]", "", st.session_state.draft_title or "draft").strip().replace(" ", "_")
            st.download_button(
                label    = "⬇ Download .docx",
                data     = docx_bytes,
                file_name= f"{safe_name}_{datetime.date.today()}.docx",
                mime     = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.caption("Write something in the draft to enable download.")

# ══════════════════════════════════════════════
#  TAB 3 — WORD TOOLS
# ══════════════════════════════════════════════
elif nav == "📖 Word Tools":
    st.markdown("### 📖 Word Tools")
    st.caption("Look up definitions, synonyms, pronunciations, or get descriptors for a sentence.")

    tool_tab = st.tabs(["📘 Define", "🔀 Synonyms", "🏷 Describe Sentence"])

    # ── Define ──
    with tool_tab[0]:
        word_in = st.text_input("Enter a word", placeholder="e.g. ephemeral, serendipity, melancholy…", key="define_input")
        if st.button("Look up", key="define_btn", use_container_width=False):
            if not word_in.strip():
                st.warning("Enter a word first.")
            else:
                with st.spinner(f"Looking up '{word_in}'…"):
                    try:
                        prompt = (
                            f'Provide a full dictionary entry for the word: "{word_in}"\n\n'
                            "Format your response EXACTLY as:\n"
                            "WORD: [word]\n"
                            "PART_OF_SPEECH: [noun/verb/adjective/etc.]\n"
                            "DEFINITION: [clear definition]\n"
                            "EXAMPLE: [one natural example sentence]\n"
                            "PRONUNCIATION: Pronounced as: [phonetic, e.g. ser-EN-i-tee]\n"
                            "ETYMOLOGY: [brief origin]\n"
                        )
                        result = call_model(prompt, max_tokens=400, temperature=0.15)
                        st.session_state.word_result = ("define", word_in, result)
                    except Exception as e:
                        st.error(str(e))

        if st.session_state.word_result and st.session_state.word_result[0] == "define":
            _, word, raw = st.session_state.word_result
            parsed = {}
            for line in raw.split("\n"):
                m = re.match(r"^([A-Z_]+):\s*(.+)$", line.strip())
                if m:
                    parsed[m.group(1)] = m.group(2).strip()

            st.markdown(f'<div class="result-card">'
                f'<div class="word-title">{parsed.get("WORD", word)}</div>'
                + (f'<span class="pos">{parsed["PART_OF_SPEECH"]}</span>' if "PART_OF_SPEECH" in parsed else "")
                + (f'<div style="color:#e8ecf4;margin-bottom:0.5rem;">{parsed["DEFINITION"]}</div>' if "DEFINITION" in parsed else "")
                + (f'<div style="font-style:italic;color:#8a93a8;font-size:0.82rem;margin-bottom:0.4rem;">"{parsed["EXAMPLE"]}"</div>' if "EXAMPLE" in parsed else "")
                + (f'<div class="pron">{parsed["PRONUNCIATION"]}</div>' if "PRONUNCIATION" in parsed else "")
                + (f'<div style="font-size:0.73rem;color:#5a6278;margin-top:0.4rem;">⌛ {parsed["ETYMOLOGY"]}</div>' if "ETYMOLOGY" in parsed else "")
                + "</div>",
                unsafe_allow_html=True,
            )

    # ── Synonyms ──
    with tool_tab[1]:
        syn_in = st.text_input("Enter a word", placeholder="e.g. happy, sad, beautiful…", key="syn_input")
        if st.button("Find synonyms", key="syn_btn"):
            if not syn_in.strip():
                st.warning("Enter a word first.")
            else:
                with st.spinner(f"Finding synonyms for '{syn_in}'…"):
                    try:
                        prompt = (
                            f'List 7 synonyms or closely related words for: "{syn_in}"\n\n'
                            "For each, use this format on a new line:\n"
                            "• [word] — [brief meaning / nuance difference]\n\n"
                            "Cover different registers: formal, informal, poetic, technical."
                        )
                        result = call_model(prompt, max_tokens=500, temperature=0.4)
                        st.session_state.word_result = ("synonyms", syn_in, result)
                    except Exception as e:
                        st.error(str(e))

        if st.session_state.word_result and st.session_state.word_result[0] == "synonyms":
            _, word, raw = st.session_state.word_result
            st.markdown(f"**Synonyms for** *{word}*")
            lines = [l for l in raw.split("\n") if l.strip().startswith("•")]
            if lines:
                for line in lines:
                    m = re.match(r"•\s*(.+?)\s*—\s*(.+)", line)
                    if m:
                        st.markdown(
                            f'<div class="result-card" style="padding:0.5rem 0.85rem;">'
                            f'<strong>{m.group(1).strip()}</strong>'
                            f'<div style="font-size:0.76rem;color:#8a93a8;margin-top:2px;">{m.group(2).strip()}</div>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )
            else:
                st.markdown(f'<div class="result-card">{raw}</div>', unsafe_allow_html=True)

    # ── Describe Sentence ──
    with tool_tab[2]:
        sent_in = st.text_area("Enter a sentence", placeholder="Paste a sentence to get mood, style, and descriptor words…", height=100, key="desc_input")
        if st.button("Describe", key="desc_btn"):
            if not sent_in.strip():
                st.warning("Enter a sentence first.")
            else:
                with st.spinner("Analysing sentence…"):
                    try:
                        prompt = (
                            f'Analyse this sentence and describe its character:\n"{sent_in}"\n\n'
                            "Respond EXACTLY as:\n"
                            "MOOD: [2-3 words for emotional tone]\n"
                            "STYLE: [2-3 words for writing style]\n"
                            "IMAGERY: [dominant image or sensory impression]\n"
                            "DESCRIPTORS: [5-6 single descriptor words, comma-separated]\n"
                            "SIMILAR_TO: [one famous author or work this resembles + one-line reason]\n"
                        )
                        result = call_model(prompt, max_tokens=400, temperature=0.5)
                        st.session_state.word_result = ("describe", sent_in[:60], result)
                    except Exception as e:
                        st.error(str(e))

        if st.session_state.word_result and st.session_state.word_result[0] == "describe":
            _, _, raw = st.session_state.word_result
            parsed = {}
            for line in raw.split("\n"):
                m = re.match(r"^([A-Z_]+):\s*(.+)$", line.strip())
                if m:
                    parsed[m.group(1)] = m.group(2).strip()

            html = '<div class="result-card">'
            for key, label, color in [
                ("MOOD",    "Mood",    "#a78bfa"),
                ("STYLE",   "Style",   "#7c6af0"),
                ("IMAGERY", "Imagery", "#d4a853"),
            ]:
                if key in parsed:
                    html += (f'<div style="margin-bottom:6px;">'
                             f'<span style="font-size:0.65rem;font-weight:600;letter-spacing:.08em;'
                             f'text-transform:uppercase;color:{color};">{label}</span><br>'
                             f'<span style="color:#c8cdd8;font-size:0.85rem;">{parsed[key]}</span></div>')

            if "DESCRIPTORS" in parsed:
                chips = "".join(
                    f'<span style="display:inline-block;font-size:0.72rem;padding:2px 10px;'
                    f'border-radius:20px;background:#1a1e2a;border:1px solid #2a2f3d;'
                    f'color:#8a93a8;margin:2px;">{d.strip()}</span>'
                    for d in parsed["DESCRIPTORS"].split(",") if d.strip()
                )
                html += f'<div style="margin:6px 0 4px;font-size:0.65rem;font-weight:600;letter-spacing:.08em;text-transform:uppercase;color:#4ade80;">Descriptors</div>{chips}'

            if "SIMILAR_TO" in parsed:
                html += (f'<div style="margin-top:8px;padding-top:6px;border-top:1px solid #2a2f3d;'
                         f'font-size:0.75rem;color:#5a6278;">📚 {parsed["SIMILAR_TO"]}</div>')

            html += "</div>"
            st.markdown(html, unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  Footer
# ─────────────────────────────────────────────
st.markdown("---")
st.markdown(
    '<div style="text-align:center;font-size:0.68rem;color:#3a3f52;padding:4px 0;">'
    'Inkwell · IBM Watsonx.ai · Llama 3.3 70B Instruct'
    '</div>',
    unsafe_allow_html=True,
)
